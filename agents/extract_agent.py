import os
import logging
import json
import requests
from typing import List, Dict, Any, Optional
from docx import Document
import fitz  # PyMuPDF
import re

logger = logging.getLogger(__name__)

class ExtractAgent:
    """Agent responsible for extracting questions from assessment documents"""
    
    def __init__(self, api_key: str = ""):
        self.api_key = api_key
        self.question_patterns = [
            r'^\d+\.\s*(.+?)(?=\n\d+\.|\n\n|$)',  # Numbered questions (20.)
            r'^\d+\.\s+\d+\s*(.+?)(?=\n\d+\.|\n\n|$)',  # Sub-questions (20.1, 20.2)
            r'^[A-Z]\.\s*(.+?)(?=\n[A-Z]\.|\n\n|$)',  # Lettered questions
            r'^Question\s+\d+:\s*(.+?)(?=\nQuestion|\n\n|$)',  # "Question X:" format
            r'^Q\d+\.\s*(.+?)(?=\nQ\d+|\n\n|$)',  # Q1, Q2 format
        ]
    
    def execute(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract questions from an assessment document using LLM-based intelligent extraction
        
        Args:
            file_path (str): Path to the assessment document
            
        Returns:
            List of dictionaries containing extracted questions with confidence scores
        """
        try:
            logger.info(f"Extracting questions from: {file_path}")
            
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Determine file type and extract text
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.docx':
                text = self._extract_from_docx(file_path)
            elif file_extension == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif file_extension == '.doc':
                text = self._extract_from_doc(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if not text or len(text.strip()) < 50:
                logger.warning("Extracted text is too short or empty")
                return []
            
            # Use LLM-based intelligent extraction if API key available
            if self.api_key:
                questions = self._extract_questions_with_llm(text)
            else:
                logger.info("No API key provided, using pattern-based extraction")
                questions = self._extract_questions_from_text(text)
            
            # Validate questions before returning
            validated_questions = self._validate_questions(questions)
            logger.info(f"Successfully extracted {len(validated_questions)} questions from {file_path}")
            return validated_questions
            
        except Exception as e:
            logger.error(f"Error extracting questions from {file_path}: {str(e)}")
            # Return empty list instead of trying to extract from undefined text
            return []

    def execute_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract questions from pasted text using LLM-based intelligent extraction
        
        Args:
            text (str): Raw text content from user paste
            
        Returns:
            List of dictionaries containing extracted questions with confidence scores
        """
        try:
            logger.info(f"Extracting questions from pasted text (length: {len(text)})")
            
            if not text or len(text.strip()) < 50:
                logger.warning("Input text is too short or empty")
                return []
            
            # Handle large texts with improved chunking
            if len(text) > 5000:  # Lower threshold for better chunking
                logger.warning(f"Large text detected ({len(text)} chars), processing in chunks")
                return self._process_in_chunks(text, chunk_size=3500)
            
            # Pre-format the text for better LLM processing
            formatted_text = self._format_assessment_text(text)
            logger.info(f"Text pre-formatted for LLM processing (length: {len(formatted_text)})")
            
            # Use LLM-based intelligent extraction if API key available
            if self.api_key:
                questions = self._extract_questions_with_llm(formatted_text)
            else:
                logger.info("No API key provided, using pattern-based extraction")
                questions = self._extract_questions_from_text(formatted_text)
            
            # Validate questions before returning
            validated_questions = self._validate_questions(questions)
            logger.info(f"Successfully extracted {len(validated_questions)} questions from pasted text")
            return validated_questions
            
        except Exception as e:
            logger.error(f"Error extracting questions from pasted text: {str(e)}")
            # Fallback to pattern-based extraction
            logger.info("Falling back to pattern-based extraction")
            try:
                return self._extract_questions_from_text(text)
            except Exception as fallback_error:
                logger.error(f"Fallback extraction also failed: {str(fallback_error)}")
                return []

    def _process_in_chunks(self, text: str, chunk_size=3500) -> List[Dict[str, Any]]:
        """
        Process large texts using improved chunking that preserves complete Q&A groups
        
        Args:
            text (str): Large text content
            chunk_size (int): Maximum size for each chunk
            
        Returns:
            List of dictionaries containing extracted questions
        """
        logger.info(f"Processing large text in chunks (max size: {chunk_size})")
        
        try:
            # Split by questions while preserving complete Q&A groups
            questions = self._split_by_questions(text)
            logger.info(f"Split text into {len(questions)} question groups")
            
            # Group into chunks that stay under token limit
            chunks = self._group_into_chunks(questions, chunk_size)
            logger.info(f"Grouped into {len(chunks)} chunks")
            
            # Process each chunk independently
            all_results = []
            for i, chunk in enumerate(chunks):
                try:
                    logger.info(f"Processing chunk {i+1}/{len(chunks)} (length: {len(chunk)})")
                    # Pre-format the chunk
                    formatted_chunk = self._format_assessment_text(chunk)
                    # Extract questions from this chunk
                    if self.api_key:
                        result = self._extract_questions_with_llm(formatted_chunk)
                    else:
                        result = self._extract_questions_from_text(formatted_chunk)
                    all_results.extend(result)
                    logger.info(f"Chunk {i+1}: Extracted {len(result)} questions")
                except Exception as e:
                    logger.error(f"Error processing chunk {i+1}: {str(e)}")
                    # Fallback to pattern extraction for this chunk
                    try:
                        result = self._extract_questions_from_text(chunk)
                        all_results.extend(result)
                        logger.info(f"Chunk {i+1}: Fallback extracted {len(result)} questions")
                    except Exception as fallback_error:
                        logger.error(f"Chunk {i+1}: Fallback also failed: {str(fallback_error)}")
                        continue
            
            # Validate all extracted questions
            validated_results = self._validate_questions(all_results)
            logger.info(f"Successfully extracted {len(validated_results)} total questions from large text")
            return validated_results
            
        except Exception as e:
            logger.error(f"Error in chunk processing: {str(e)}")
            return []

    def _split_by_questions(self, text: str) -> List[str]:
        """
        Split text on question numbers but keep each question with its choices
        
        Args:
            text (str): Raw text content
            
        Returns:
            List of question groups
        """
        try:
            # Split on question numbers (more robust pattern)
            questions = re.split(r'(?=\d+\.)', text)
            
            # Clean up empty questions and ensure proper formatting
            cleaned_questions = []
            for question in questions:
                question = question.strip()
                if question and len(question) > 10:  # Only keep substantial questions
                    cleaned_questions.append(question)
            
            return cleaned_questions
        except Exception as e:
            logger.error(f"Error splitting text by questions: {str(e)}")
            return [text]  # Return original text as single chunk

    def _group_into_chunks(self, questions: List[str], max_size: int) -> List[str]:
        """
        Group questions into chunks that stay under the token limit
        
        Args:
            questions (List[str]): List of question groups
            max_size (int): Maximum size for each chunk
            
        Returns:
            List of chunks
        """
        try:
            chunks = []
            current = ""
            
            for q in questions:
                # If adding this question would exceed the limit, start a new chunk
                if len(current + q) > max_size and current:
                    chunks.append(current.strip())
                    current = q
                else:
                    # Add to current chunk
                    if current:
                        current += "\n" + q
                    else:
                        current = q
            
            # Add the last chunk
            if current:
                chunks.append(current.strip())
            
            return chunks
        except Exception as e:
            logger.error(f"Error grouping into chunks: {str(e)}")
            return ["\n".join(questions)]  # Return all questions as single chunk

    def _format_assessment_text(self, text: str) -> str:
        """
        Pre-format assessment text for better LLM processing
        
        Args:
            text (str): Raw assessment text
            
        Returns:
            str: Formatted text with consistent structure
        """
        try:
            lines = text.split('\n')
            formatted = ["# Assessment Marking Guide\n"]
            current_q = 0
            in_choices = False
            choice_count = 0
            
            for line in lines:
                line = line.strip()
                if not line or line.startswith(('PC-', 'KE-')):
                    continue
                    
                if re.match(r'^\d+\.', line):
                    # Extract the original question number
                    question_match = re.match(r'^(\d+)\.', line)
                    if question_match:
                        original_number = question_match.group(1)
                        formatted.append(f"\n## Question {original_number}")
                        formatted.append(f"**{line[line.find('.')+1:].strip()}**\n")
                    in_choices = False
                    choice_count = 0
                elif re.match(r'^\d+\.\s*\d+', line):
                    part = re.search(r'(\d+\.\s*\d+)', line).group(1)
                    formatted.append(f"### Part {part} - {line[len(part):].strip()}")
                    in_choices = False
                    choice_count = 0
                elif line and not line.startswith(('Select', 'Choose')) and len(line) > 10:
                    # This might be a choice
                    choice_count += 1
                    letter = chr(64 + choice_count)  # A, B, C, D...
                    formatted.append(f"- {letter}) {line}")
            
            return '\n'.join(formatted)
        except Exception as e:
            logger.error(f"Error formatting assessment text: {str(e)}")
            return text  # Return original text if formatting fails
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file including tables"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(file_path)
            text = []
            
            for page in doc:
                text.append(page.get_text())
            
            doc.close()
            return '\n'.join(text)
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            raise
    
    def _extract_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file (basic implementation)"""
        try:
            # For DOC files, we'll use a basic text extraction
            # In production, you might want to use antiword or similar tools
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"Error extracting from DOC: {str(e)}")
            raise
    
    def _extract_questions_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract questions from text using various patterns with hierarchical support"""
        try:
            questions = []
            lines = text.split('\n')
            question_counter = 1
            used_ids = set()
            
            # First pass: identify main questions and their sub-questions
            main_questions = {}
            current_main = None
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Check for question heading/topic pattern (e.g., "5. Communication strategies...")
                main_match = re.match(r'^(\d+)\.\s*(.+?)$', line)
                if main_match:
                    question_num = main_match.group(1)
                    question_text = main_match.group(2).strip()
                    if len(question_text) > 5:  # Minimum title length
                        question_id = f"Q{question_num}"
                        current_main = {
                            'id': question_id,
                            'question_id': question_id,  # Ensure both id and question_id are set
                            'text': question_text,
                            'question_number': question_num,
                            'line_number': i + 1,
                            'pattern_used': 'question_heading',
                            'type': 'heading',
                            'question_type': 'topic',
                            'sub_questions': []
                        }
                        main_questions[question_num] = current_main
                        # Don't add headings as standalone questions - only as containers
                        continue
                
                # Check for sub-question pattern (e.g., "20.1", "20.2")
                sub_match = re.match(r'^(\d+)\.\s*(\d+)\s*(.+?)$', line)
                if sub_match:
                    main_num = sub_match.group(1)
                    sub_num = sub_match.group(2)
                    sub_text = sub_match.group(3).strip()
                    
                    if len(sub_text) > 10:  # Minimum sub-question length
                        question_type = self._classify_question_type(sub_text)
                        # Generate unique sub-question ID
                        sub_question_id = f"Q{main_num}.{sub_num}"
                        if sub_question_id in used_ids:
                            # If ID already exists, add a suffix
                            counter = 1
                            while f"{sub_question_id}_{counter}" in used_ids:
                                counter += 1
                            sub_question_id = f"{sub_question_id}_{counter}"
                        used_ids.add(sub_question_id)
                        sub_question = {
                            'id': sub_question_id,
                            'question_id': sub_question_id,  # Ensure both id and question_id are set
                            'text': sub_text,
                            'question_number': f"{main_num}.{sub_num}",
                            'line_number': i + 1,
                            'pattern_used': 'sub_question',
                            'type': question_type,
                            'question_type': question_type,
                            'parent_question': f"Q{main_num}",
                            'choices': self._extract_choices(sub_text, question_type) if question_type == 'mcq' else []
                        }
                        
                        # Add to main question's sub_questions if it exists
                        if main_num in main_questions:
                            main_questions[main_num]['sub_questions'].append(sub_question)
                        else:
                            # If main question doesn't exist, create it
                            main_question_id = f"Q{main_num}"
                            current_main = {
                                'id': main_question_id,
                                'question_id': main_question_id,  # Ensure both id and question_id are set
                                'text': f"Question {main_num}",
                                'question_number': main_num,
                                'line_number': i + 1,
                                'pattern_used': 'main_question_placeholder',
                                'type': 'unknown',
                                'sub_questions': [sub_question]
                            }
                            main_questions[main_num] = current_main
                            questions.append(current_main)
                        
                        questions.append(sub_question)
                        continue
                
                # Try other question patterns
                for pattern in self.question_patterns:
                    try:
                        match = re.match(pattern, line, re.IGNORECASE)
                        if match:
                            question_text = match.group(1).strip()
                            if len(question_text) > 10:  # Minimum question length
                                # Get the full question text including choices (next few lines)
                                full_question_text = self._get_full_question_text(lines, i)
                                question_type = self._classify_question_type(full_question_text)
                                
                                # Generate unique question ID
                                while f"Q{question_counter}" in used_ids:
                                    question_counter += 1
                                question_id = f"Q{question_counter}"
                                used_ids.add(question_id)
                                question_counter += 1
                                questions.append({
                                    'id': question_id,
                                    'question_id': question_id,  # Ensure both id and question_id are set
                                    'text': full_question_text,
                                    'line_number': i + 1,
                                    'pattern_used': pattern,
                                    'type': question_type,
                                    'question_type': question_type,
                                    'choices': self._extract_choices(full_question_text, question_type)
                                })
                                break
                    except Exception as pattern_error:
                        logger.debug(f"Pattern matching error for line {i+1}: {str(pattern_error)}")
                        continue
            
            # If no structured questions found, try to extract potential questions
            if not questions:
                questions = self._extract_potential_questions(text)
            
            # Validate and ensure all questions have required fields
            validated_questions = self._validate_questions(questions)
            
            return validated_questions
            
        except Exception as e:
            logger.error(f"Error extracting questions from text: {str(e)}")
            return []
    
    def _extract_questions_with_llm(self, text: str) -> List[Dict[str, Any]]:
        """Use LLM to intelligently extract and validate questions"""
        try:
            # Truncate text if too long for API
            if len(text) > 6000:  # Reduced from 8000 to be safer
                text = text[:6000] + "\n[TRUNCATED]"
                logger.warning(f"Text truncated to {len(text)} characters for LLM processing")
            
            # Create simplified prompt to avoid JSON parsing issues
            prompt = f"""Extract questions from this VET assessment document. Return ONLY valid JSON.

Text:
{text}

Return this JSON structure ONLY (no markdown, no other text):

{{
  "questions": [
    {{
      "id": "Q1",
      "text": "Question text here",
      "question_number": "1",
      "type": "mcq",
      "question_type": "mcq",
      "choices": ["Choice 1", "Choice 2"],
      "confidence": "high"
    }}
  ]
}}

Focus on numbered questions and their choices. Question types: mcq, short_answer, essay, scenario, practical."""
            
            # Call LLM API
            response = self._call_llm_api(prompt)
            
            if response:
                try:
                    # Clean up response - remove markdown code blocks if present
                    cleaned_response = response.strip()
                    if cleaned_response.startswith('```json'):
                        cleaned_response = cleaned_response[7:]
                    if cleaned_response.endswith('```'):
                        cleaned_response = cleaned_response[:-3]
                    cleaned_response = cleaned_response.strip()
                    
                    # Extract JSON from response
                    json_start = cleaned_response.find('{')
                    json_end = cleaned_response.rfind('}') + 1
                    
                    if json_start != -1 and json_end > json_start:
                        json_str = cleaned_response[json_start:json_end]
                        data = json.loads(json_str)
                        questions = data.get('questions', [])
                        
                        # Convert to standard format
                        formatted_questions = []
                        for i, q in enumerate(questions):
                            question_id = q.get('id', f"Q{i+1}")
                            formatted_questions.append({
                                'id': question_id,
                                'question_id': question_id,  # Ensure both id and question_id are set
                                'text': q.get('text', ''),
                                'question_number': q.get('question_number', str(i+1)),
                                'type': q.get('type', 'unknown'),
                                'question_type': q.get('question_type', 'unknown'),
                                'choices': q.get('choices', []),
                                'confidence': q.get('confidence', 'medium'),
                                'line_number': i + 1,
                                'pattern_used': 'llm_extraction'
                            })
                        
                        # Validate the formatted questions
                        validated_questions = self._validate_questions(formatted_questions)
                        return validated_questions
                    else:
                        logger.warning("No valid JSON found in LLM response")
                        return self._extract_questions_from_text(text)
                        
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse LLM response as JSON: {e}")
                    logger.debug(f"Response was: {response[:500]}...")
                    return self._extract_questions_from_text(text)
            else:
                logger.warning("LLM API call failed, falling back to pattern extraction")
                return self._extract_questions_from_text(text)
                
        except Exception as e:
            logger.error(f"Error in LLM extraction: {str(e)}")
            return self._extract_questions_from_text(text)
    
    def _call_llm_api(self, prompt: str) -> Optional[str]:
        """Call LLM API for intelligent question extraction"""
        try:
            if not self.api_key:
                logger.warning("No API key available for LLM extraction")
                return None
            
            headers = {
                'Content-Type': 'application/json',
            }
            
            data = {
                "contents": [{
                    "parts": [{
                        "text": prompt
                    }]
                }],
                "generationConfig": {
                    "temperature": 0.1,
                    "maxOutputTokens": 2000,
                    "topP": 0.8,
                    "topK": 40
                }
            }
            
            response = requests.post(
                "https://generativelanguage.googleapis.com/v1/models/gemini-2.0-flash:generateContent",
                headers=headers,
                params={'key': self.api_key},
                json=data,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'candidates' in result and len(result['candidates']) > 0:
                    content = result['candidates'][0].get('content', {})
                    parts = content.get('parts', [])
                    if parts and 'text' in parts[0]:
                        return parts[0]['text']
                return None
            else:
                logger.error(f"LLM API error: {response.status_code} - {response.text}")
                return None
                
        except Exception as e:
            logger.error(f"LLM API error: {str(e)}")
            return None
    
    def _extract_choices(self, question_text: str, question_type: str) -> List[str]:
        """Extract MCQ choices from question text with enhanced pattern matching"""
        if question_type != 'mcq':
            return []
        
        try:
            choices = []
            lines = question_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Enhanced choice patterns
                choice_patterns = [
                    r'^[A-D]\.\s*(.+)$',  # A. Choice text
                    r'^[1-4]\.\s*(.+)$',  # 1. Choice text
                    r'^[-•]\s*(.+)$',     # - Choice text
                    r'^[a-d]\)\s*(.+)$',  # a) Choice text
                    r'^[1-4]\)\s*(.+)$',  # 1) Choice text
                    r'^[A-D]\)\s*(.+)$',  # A) Choice text
                    r'^[a-d]\.\s*(.+)$',  # a. Choice text
                    r'^\([A-D]\)\s*(.+)$', # (A) Choice text
                    r'^\([1-4]\)\s*(.+)$', # (1) Choice text
                ]
                
                for pattern in choice_patterns:
                    match = re.match(pattern, line, re.IGNORECASE)
                    if match:
                        choice_text = match.group(1).strip()
                        if len(choice_text) > 2:  # Minimum choice length
                            # Clean up choice text (remove any duplicate indicators)
                            choice_text = re.sub(r'^[A-D]\.\s*[A-D]\.\s*', 'A. ', choice_text)
                            choice_text = re.sub(r'^[1-4]\.\s*[1-4]\.\s*', '1. ', choice_text)
                            choices.append(choice_text)
                        break
            
            return choices
        except Exception as e:
            logger.error(f"Error extracting choices: {str(e)}")
            return []

    def _extract_potential_questions(self, text: str) -> List[Dict[str, Any]]:
        """Extract potential questions when structured patterns aren't found"""
        try:
            questions = []
            lines = text.split('\n')
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # Look for lines that end with question marks or contain question words
                if (line.endswith('?') or 
                    any(word in line.lower() for word in ['what', 'how', 'why', 'when', 'where', 'which', 'who'])):
                    
                    question_id = f"Q{len(questions) + 1}"
                    questions.append({
                        'id': question_id,
                        'question_id': question_id,  # Ensure both id and question_id are set
                        'text': line,
                        'question_number': str(len(questions) + 1),
                        'line_number': i + 1,
                        'pattern_used': 'question_detection',
                        'type': self._classify_question_type(line),
                        'question_type': self._classify_question_type(line),
                        'choices': []
                    })
            
            return questions
        except Exception as e:
            logger.error(f"Error extracting potential questions: {str(e)}")
            return []
    
    def _get_full_question_text(self, lines: List[str], start_line: int) -> str:
        """Get the full question text including choices from the given line onwards"""
        try:
            question_lines = []
            i = start_line
            
            # Add the main question line
            question_lines.append(lines[i].strip())
            i += 1
            
            # Continue adding lines until we hit another question or empty line
            while i < len(lines):
                line = lines[i].strip()
                
                # Stop if we hit another question pattern
                if (re.match(r'^\d+\.\s*\d+\s*\w+', line) or  # Sub-question
                    re.match(r'^\d+\.\s*\w+', line) or          # Main question
                    re.match(r'^[A-Z]\.\s*\w+', line) or        # Lettered question
                    re.match(r'^Question\s+\d+:', line) or      # "Question X:" format
                    re.match(r'^Q\d+\.\s*\w+', line)):          # Q1, Q2 format
                    break
                
                # Stop if we hit an empty line followed by another question
                if not line and i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if (re.match(r'^\d+\.\s*\w+', next_line) or
                        re.match(r'^Question\s+\d+:', next_line) or
                        re.match(r'^Q\d+\.\s*\w+', next_line)):
                        break
                
                # Add the line if it's not empty or if it's part of the question
                if line or (question_lines and question_lines[-1]):  # Don't add empty lines unless they're between content
                    question_lines.append(line)
                
                i += 1
            
            return '\n'.join(question_lines).strip()
            
        except Exception as e:
            logger.error(f"Error getting full question text: {str(e)}")
            return lines[start_line].strip()
    
    def _classify_question_type(self, question_text: str) -> str:
        """Classify the type of question based on content with enhanced MCQ detection"""
        try:
            text_lower = question_text.lower()
            
            # Enhanced MCQ detection - look for choice patterns in the text
            choice_patterns = [
                r'[A-D]\.\s*\w+',  # A. Choice
                r'[1-4]\.\s*\w+',  # 1. Choice
                r'[a-d]\)\s*\w+',  # a) Choice
                r'[1-4]\)\s*\w+',  # 1) Choice
                r'[A-D]\)\s*\w+',  # A) Choice
                r'\([A-D]\)\s*\w+', # (A) Choice
                r'\([1-4]\)\s*\w+', # (1) Choice
            ]
            
            # Check if text contains choice patterns
            for pattern in choice_patterns:
                if re.search(pattern, question_text, re.IGNORECASE):
                    return 'mcq'
            
            # Check for explicit MCQ indicators
            if any(word in text_lower for word in ['multiple choice', 'select', 'choose', 'mcq', 'select the best', 'choose the correct']):
                return 'mcq'
            elif any(word in text_lower for word in ['true', 'false', 't/f']):
                return 'true_false'
            elif any(word in text_lower for word in ['describe', 'explain', 'discuss', 'analyze', 'elaborate']):
                return 'essay'
            elif any(word in text_lower for word in ['list', 'name', 'identify', 'state']):
                return 'short_answer'
            elif any(word in text_lower for word in ['how', 'what', 'why', 'when', 'where', 'which', 'who']) and '?' in question_text:
                return 'short_answer'
            elif any(word in text_lower for word in ['scenario', 'case study', 'situation']):
                return 'scenario'
            elif any(word in text_lower for word in ['demonstrate', 'show', 'perform', 'practice']):
                return 'practical'
            elif '?' in question_text:
                return 'question'
            else:
                return 'unknown'
        except Exception as e:
            logger.error(f"Error classifying question type: {str(e)}")
            return 'unknown'
    
    def _validate_questions(self, questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Validate and ensure all questions have required fields"""
        try:
            validated_questions = []
            for i, question in enumerate(questions):
                # Ensure question_id is present (mapping agent requires this)
                if 'question_id' not in question:
                    question['question_id'] = question.get('id', f"Q{i+1}")
                
                # Ensure id is present
                if 'id' not in question:
                    question['id'] = question['question_id']
                
                # Ensure text is present
                if 'text' not in question or not question['text']:
                    question['text'] = f"Question {i+1}"
                
                # Ensure question_number is present
                if 'question_number' not in question:
                    question['question_number'] = str(i+1)
                
                # Ensure type and question_type are present
                if 'type' not in question:
                    question['type'] = 'unknown'
                if 'question_type' not in question:
                    question['question_type'] = question['type']
                
                # Ensure choices is present
                if 'choices' not in question:
                    question['choices'] = []
                
                # Ensure confidence is present
                if 'confidence' not in question:
                    question['confidence'] = 'medium'
                
                # Ensure line_number is present
                if 'line_number' not in question:
                    question['line_number'] = i + 1
                
                # Ensure pattern_used is present
                if 'pattern_used' not in question:
                    question['pattern_used'] = 'validated'
                
                validated_questions.append(question)
            
            return validated_questions
            
        except Exception as e:
            logger.error(f"Error validating questions: {str(e)}")
            return questions

    def get_question_statistics(self, questions: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get statistics about extracted questions"""
        try:
            if not questions:
                return {
                    'total_questions': 0,
                    'question_types': {},
                    'patterns_used': {}
                }
            
            question_types = {}
            patterns_used = {}
            
            for question in questions:
                q_type = question.get('type', 'unknown')
                pattern = question.get('pattern_used', 'unknown')
                
                question_types[q_type] = question_types.get(q_type, 0) + 1
                patterns_used[pattern] = patterns_used.get(pattern, 0) + 1
            
            return {
                'total_questions': len(questions),
                'question_types': question_types,
                'patterns_used': patterns_used
            }
        except Exception as e:
            logger.error(f"Error calculating question statistics: {str(e)}")
            return {
                'total_questions': 0,
                'question_types': {},
                'patterns_used': {}
            } 